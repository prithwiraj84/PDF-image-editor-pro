import sys
import os
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QLabel, QFileDialog, 
                             QTextEdit, QMessageBox, QScrollArea, QGraphicsView, 
                             QGraphicsScene, QGraphicsTextItem, QGraphicsRectItem,
                             QColorDialog, QFontComboBox, QSpinBox, QComboBox,
                             QToolBar, QAction, QStatusBar, QSplitter, QGroupBox,
                             QSlider, QCheckBox)
from PyQt5.QtCore import Qt, QPointF, QRectF, pyqtSignal, QSize
from PyQt5.QtGui import QFont, QColor, QPixmap, QImage, QPainter, QPen, QBrush
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import io
from docx import Document
from docx.shared import Pt, RGBColor
import json
from typing import List, Dict, Tuple, Optional
import copy

class PDFCanvas(QGraphicsView):
    """Custom canvas for displaying and editing PDF pages"""
    textSelected = pyqtSignal(float, float, str, dict)  # x, y, text, font_info
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.scene = QGraphicsScene(self)
        self.setScene(self.scene)
        self.setDragMode(QGraphicsView.RubberBandDrag)
        self.setRenderHint(QPainter.Antialiasing)
        self.setRenderHint(QPainter.SmoothPixmapTransform)
        
        self.pdf_doc = None
        self.current_page = 0
        self.page_pixmap = None
        self.text_blocks = []  # List of text blocks with positions
        self.selected_text_item = None
        self.edit_mode = False
        self.zoom_factor = 1.0
        
    def load_pdf(self, file_path: str):
        """Load PDF document"""
        try:
            self.pdf_doc = fitz.open(file_path)
            self.current_page = 0
            self.display_page(0)
            return True
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load PDF: {str(e)}")
            return False
    
    def display_page(self, page_num: int):
        """Display a specific page of the PDF"""
        if not self.pdf_doc or page_num < 0 or page_num >= len(self.pdf_doc):
            return
        
        self.current_page = page_num
        page = self.pdf_doc[page_num]
        
        # Render page to pixmap
        mat = fitz.Matrix(2, 2)  # Zoom factor
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")
        
        image = QImage()
        image.loadFromData(img_data)
        self.page_pixmap = QPixmap.fromImage(image)
        
        # Clear scene
        self.scene.clear()
        self.text_blocks = []
        
        # Add page image
        self.scene.addPixmap(self.page_pixmap)
        
        # Extract text blocks with positions
        self.extract_text_blocks(page)
        
        # Add text blocks as editable items
        for block in self.text_blocks:
            self.add_text_item(block)
    
    def extract_text_blocks(self, page):
        """Extract text blocks with their positions and font information"""
        text_dict = page.get_text("dict")
        
        for block in text_dict.get("blocks", []):
            if "lines" in block:  # Text block
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            bbox = span["bbox"]
                            font_info = {
                                "font": span.get("font", "Arial"),
                                "size": span.get("size", 12),
                                "flags": span.get("flags", 0),
                                "color": span.get("color", 0)
                            }
                            
                            self.text_blocks.append({
                                "text": text,
                                "x": bbox[0],
                                "y": bbox[1],
                                "width": bbox[2] - bbox[0],
                                "height": bbox[3] - bbox[1],
                                "font_info": font_info
                            })
    
    def add_text_item(self, block: dict):
        """Add a text block as an editable item on the canvas"""
        # Scale coordinates (PDF coordinates to screen coordinates)
        scale = 2.0  # Match the matrix scale
        x = block["x"] * scale
        y = block["y"] * scale
        
        # Create text item
        text_item = QGraphicsTextItem(block["text"])
        font = QFont(block["font_info"]["font"], int(block["font_info"]["size"]))
        text_item.setFont(font)
        text_item.setPos(x, y)
        text_item.setFlag(QGraphicsTextItem.ItemIsMovable, True)
        text_item.setFlag(QGraphicsTextItem.ItemIsSelectable, True)
        
        # Store original data
        text_item.setData(0, block)
        
        self.scene.addItem(text_item)
    
    def mousePressEvent(self, event):
        """Handle mouse clicks for text selection"""
        if event.button() == Qt.LeftButton and self.edit_mode:
            item = self.itemAt(event.pos())
            if isinstance(item, QGraphicsTextItem):
                self.selected_text_item = item
                # Highlight selected item
                for item in self.scene.items():
                    if isinstance(item, QGraphicsTextItem):
                        item.setDefaultTextColor(QColor(0, 0, 0))
                item.setDefaultTextColor(QColor(0, 100, 200))
                
                block = item.data(0)
                if block:
                    self.textSelected.emit(
                        block["x"], block["y"], 
                        block["text"], block["font_info"]
                    )
        
        super().mousePressEvent(event)
    
    def wheelEvent(self, event):
        """Handle mouse wheel for zooming"""
        if event.modifiers() & Qt.ControlModifier:
            # Zoom with Ctrl + Wheel
            delta = event.angleDelta().y()
            if delta > 0:
                self.scale(1.1, 1.1)
                self.zoom_factor *= 1.1
            else:
                self.scale(0.9, 0.9)
                self.zoom_factor *= 0.9
        else:
            super().wheelEvent(event)
    
    def get_page_image(self):
        """Get current page as QImage"""
        return self.page_pixmap.toImage()

class PhotoEditor(QWidget):
    """Widget for editing photos with OCR"""
    textRegionSelected = pyqtSignal(dict)  # Emit selected text region
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.image = None
        self.text_regions = []
        self.selected_region = None
        self.layout = QVBoxLayout(self)
        
        self.scroll_area = QScrollArea()
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.mousePressEvent = self.on_image_click
        self.scroll_area.setWidget(self.image_label)
        self.scroll_area.setWidgetResizable(True)
        
        self.layout.addWidget(self.scroll_area)
        
        # Add OCR button
        ocr_btn = QPushButton("🔍 Extract Text (OCR)")
        ocr_btn.clicked.connect(self.extract_text_with_ocr)
        self.layout.addWidget(ocr_btn)
    
    def on_image_click(self, event):
        """Handle clicks on image to select text regions"""
        if self.text_regions:
            # Find clicked region
            click_x = event.pos().x()
            click_y = event.pos().y()
            
            for region in self.text_regions:
                if (region["x"] <= click_x <= region["x"] + region["width"] and
                    region["y"] <= click_y <= region["y"] + region["height"]):
                    self.selected_region = region
                    self.textRegionSelected.emit(region)
                    self.highlight_region(region)
                    break
    
    def load_image(self, file_path: str):
        """Load image file"""
        try:
            self.image = Image.open(file_path)
            self.display_image()
            return True
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load image: {str(e)}")
            return False
    
    def display_image(self):
        """Display the current image"""
        if self.image:
            # Convert PIL to QPixmap
            img_byte_arr = io.BytesIO()
            self.image.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            
            pixmap = QPixmap()
            pixmap.loadFromData(img_byte_arr.read())
            
            self.image_label.setPixmap(pixmap)
    
    def extract_text_with_ocr(self):
        """Extract text from image using OCR"""
        if not self.image:
            QMessageBox.warning(self, "No Image", "Please load an image first.")
            return []
        
        try:
            # Convert to RGB if needed
            if self.image.mode != 'RGB':
                self.image = self.image.convert('RGB')
            
            # Perform OCR with detailed output
            ocr_data = pytesseract.image_to_data(
                self.image, 
                output_type=pytesseract.Output.DICT,
                config='--psm 6'  # Assume uniform block of text
            )
            
            text_regions = []
            n_boxes = len(ocr_data['text'])
            
            for i in range(n_boxes):
                text = ocr_data['text'][i].strip()
                conf = int(ocr_data['conf'][i]) if ocr_data['conf'][i] != '-1' else 0
                
                if text and conf > 30:  # Confidence threshold
                    x = ocr_data['left'][i]
                    y = ocr_data['top'][i]
                    w = ocr_data['width'][i]
                    h = ocr_data['height'][i]
                    
                    text_regions.append({
                        "text": text,
                        "x": x,
                        "y": y,
                        "width": w,
                        "height": h,
                        "confidence": conf,
                        "font_info": {
                            "font": "Arial",  # OCR doesn't provide font info
                            "size": max(h * 0.75, 10),  # Estimate size, minimum 10
                            "color": (0, 0, 0)
                        }
                    })
            
            self.text_regions = text_regions
            
            # Update display with highlighted regions
            self.display_image_with_regions()
            
            return text_regions
        except Exception as e:
            error_msg = str(e)
            if "tesseract" in error_msg.lower():
                QMessageBox.warning(
                    self, "OCR Error", 
                    f"Tesseract OCR not found.\n\n"
                    f"Please install Tesseract:\n"
                    f"Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki\n"
                    f"macOS: brew install tesseract\n"
                    f"Linux: sudo apt-get install tesseract-ocr"
                )
            else:
                QMessageBox.warning(self, "OCR Error", f"OCR failed: {error_msg}")
            return []
    
    def display_image_with_regions(self):
        """Display image with highlighted text regions"""
        if not self.image:
            return
        
        # Create a copy for display
        display_img = self.image.copy()
        draw = ImageDraw.Draw(display_img)
        
        # Draw rectangles around text regions
        for region in self.text_regions:
            x, y = region["x"], region["y"]
            w, h = region["width"], region["height"]
            draw.rectangle(
                [x, y, x + w, y + h],
                outline=(255, 0, 0),
                width=2
            )
        
        # Convert to QPixmap
        img_byte_arr = io.BytesIO()
        display_img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        
        pixmap = QPixmap()
        pixmap.loadFromData(img_byte_arr.read())
        self.image_label.setPixmap(pixmap)
    
    def highlight_region(self, region):
        """Highlight a specific text region"""
        if not self.image:
            return
        
        display_img = self.image.copy()
        draw = ImageDraw.Draw(display_img)
        
        # Highlight selected region
        x, y = region["x"], region["y"]
        w, h = region["width"], region["height"]
        draw.rectangle(
            [x, y, x + w, y + h],
            outline=(0, 255, 0),
            width=3
        )
        
        # Convert to QPixmap
        img_byte_arr = io.BytesIO()
        display_img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        
        pixmap = QPixmap()
        pixmap.loadFromData(img_byte_arr.read())
        self.image_label.setPixmap(pixmap)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF & Photo Text Editor")
        self.setGeometry(100, 100, 1400, 900)
        
        # Data storage
        self.pdf_path = None
        self.photo_path = None
        self.edit_history = []
        self.history_index = -1
        self.current_edits = {}  # Store current edits
        
        # Initialize UI
        self.init_ui()
        
        # Set Tesseract path (Windows default)
        if sys.platform == 'win32':
            tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            if os.path.exists(tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
    
    def init_ui(self):
        """Initialize the user interface"""
        # Create central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # Create toolbar
        self.create_toolbar()
        
        # Create splitter for main content
        splitter = QSplitter(Qt.Horizontal)
        
        # Left panel - Editor controls
        left_panel = self.create_editor_panel()
        splitter.addWidget(left_panel)
        
        # Right panel - Document viewer
        right_panel = self.create_viewer_panel()
        splitter.addWidget(right_panel)
        
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 2)
        
        main_layout.addWidget(splitter)
        
        # Create status bar
        self.statusBar = QStatusBar()
        self.setStatusBar(self.statusBar)
        self.statusBar.showMessage("Ready")
    
    def create_toolbar(self):
        """Create the main toolbar"""
        toolbar = QToolBar("Main Toolbar")
        toolbar.setStyleSheet("""
            QToolBar {
                spacing: 5px;
                padding: 5px;
            }
            QToolButton {
                padding: 5px;
                border-radius: 3px;
            }
            QToolButton:hover {
                background-color: #e0e0e0;
            }
        """)
        self.addToolBar(toolbar)
        
        # File operations
        open_pdf_action = QAction("📄 Open PDF", self)
        open_pdf_action.setToolTip("Open a PDF document")
        open_pdf_action.triggered.connect(self.open_pdf)
        toolbar.addAction(open_pdf_action)
        
        open_photo_action = QAction("🖼️ Open Photo", self)
        open_photo_action.setToolTip("Open an image file")
        open_photo_action.triggered.connect(self.open_photo)
        toolbar.addAction(open_photo_action)
        
        toolbar.addSeparator()
        
        # Edit mode toggle
        self.edit_mode_action = QAction("✏️ Edit Mode", self)
        self.edit_mode_action.setCheckable(True)
        self.edit_mode_action.setToolTip("Enable/disable edit mode (Click text to edit)")
        self.edit_mode_action.triggered.connect(self.toggle_edit_mode)
        toolbar.addAction(self.edit_mode_action)
        
        toolbar.addSeparator()
        
        # Undo/Redo
        undo_action = QAction("↶ Undo", self)
        undo_action.setShortcut("Ctrl+Z")
        undo_action.setToolTip("Undo last action (Ctrl+Z)")
        undo_action.triggered.connect(self.undo)
        toolbar.addAction(undo_action)
        
        redo_action = QAction("↷ Redo", self)
        redo_action.setShortcut("Ctrl+Y")
        redo_action.setToolTip("Redo last action (Ctrl+Y)")
        redo_action.triggered.connect(self.redo)
        toolbar.addAction(redo_action)
        
        toolbar.addSeparator()
        
        # Export
        export_pdf_action = QAction("💾 Save as PDF", self)
        export_pdf_action.setToolTip("Export document as PDF")
        export_pdf_action.triggered.connect(self.export_pdf)
        toolbar.addAction(export_pdf_action)
        
        export_docx_action = QAction("📝 Save as DOCX", self)
        export_docx_action.setToolTip("Export document as Word (DOCX)")
        export_docx_action.triggered.connect(self.export_docx)
        toolbar.addAction(export_docx_action)
        
        export_png_action = QAction("🖼️ Save as PNG", self)
        export_png_action.setToolTip("Export current page/image as PNG")
        export_png_action.triggered.connect(self.export_png)
        toolbar.addAction(export_png_action)
    
    def create_editor_panel(self):
        """Create the left editor panel"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # Text editor group
        editor_group = QGroupBox("Text Editor")
        editor_layout = QVBoxLayout()
        
        # Text input
        self.text_input = QTextEdit()
        self.text_input.setPlaceholderText("Selected text will appear here. Click on text in the document to edit.")
        editor_layout.addWidget(QLabel("Text Content:"))
        editor_layout.addWidget(self.text_input)
        
        # Font controls
        font_group = QGroupBox("Font Settings")
        font_layout = QVBoxLayout()
        
        font_layout.addWidget(QLabel("Font Family:"))
        self.font_combo = QFontComboBox()
        font_layout.addWidget(self.font_combo)
        
        size_layout = QHBoxLayout()
        size_layout.addWidget(QLabel("Size:"))
        self.font_size = QSpinBox()
        self.font_size.setRange(6, 144)
        self.font_size.setValue(12)
        size_layout.addWidget(self.font_size)
        font_layout.addLayout(size_layout)
        
        # Color picker
        color_layout = QHBoxLayout()
        color_layout.addWidget(QLabel("Color:"))
        self.color_btn = QPushButton("Choose Color")
        self.color_btn.clicked.connect(self.choose_color)
        self.current_color = QColor(0, 0, 0)
        color_layout.addWidget(self.color_btn)
        font_layout.addLayout(color_layout)
        
        font_group.setLayout(font_layout)
        editor_layout.addWidget(font_group)
        
        # Apply button
        self.apply_btn = QPushButton("Apply Changes")
        self.apply_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:pressed {
                background-color: #3d8b40;
            }
        """)
        self.apply_btn.clicked.connect(self.apply_text_changes)
        editor_layout.addWidget(self.apply_btn)
        
        editor_group.setLayout(editor_layout)
        layout.addWidget(editor_group)
        
        # Page navigation (for PDF)
        nav_group = QGroupBox("Navigation")
        nav_layout = QVBoxLayout()
        
        page_nav_layout = QHBoxLayout()
        self.prev_btn = QPushButton("◀ Previous")
        self.prev_btn.clicked.connect(self.prev_page)
        page_nav_layout.addWidget(self.prev_btn)
        
        self.page_label = QLabel("Page 1 of 1")
        page_nav_layout.addWidget(self.page_label)
        
        self.next_btn = QPushButton("Next ▶")
        self.next_btn.clicked.connect(self.next_page)
        page_nav_layout.addWidget(self.next_btn)
        nav_layout.addLayout(page_nav_layout)
        
        # Zoom controls
        zoom_layout = QHBoxLayout()
        zoom_layout.addWidget(QLabel("Zoom:"))
        self.zoom_slider = QSlider(Qt.Horizontal)
        self.zoom_slider.setMinimum(50)
        self.zoom_slider.setMaximum(300)
        self.zoom_slider.setValue(100)
        self.zoom_slider.valueChanged.connect(self.on_zoom_changed)
        zoom_layout.addWidget(self.zoom_slider)
        self.zoom_label = QLabel("100%")
        zoom_layout.addWidget(self.zoom_label)
        nav_layout.addLayout(zoom_layout)
        
        nav_group.setLayout(nav_layout)
        layout.addWidget(nav_group)
        
        # Additional features
        features_group = QGroupBox("Features")
        features_layout = QVBoxLayout()
        
        self.show_text_blocks = QCheckBox("Show Text Blocks")
        self.show_text_blocks.setChecked(True)
        features_layout.addWidget(self.show_text_blocks)
        
        self.auto_detect_font = QCheckBox("Auto-detect Font")
        self.auto_detect_font.setChecked(True)
        features_layout.addWidget(self.auto_detect_font)
        
        features_group.setLayout(features_layout)
        layout.addWidget(features_group)
        
        layout.addStretch()
        
        return panel
    
    def create_viewer_panel(self):
        """Create the right viewer panel"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # Tab widget for PDF and Photo
        from PyQt5.QtWidgets import QTabWidget
        self.tabs = QTabWidget()
        
        # PDF tab
        self.pdf_canvas = PDFCanvas()
        self.pdf_canvas.textSelected.connect(self.on_text_selected)
        self.tabs.addTab(self.pdf_canvas, "PDF Viewer")
        
        # Photo tab
        photo_widget = QWidget()
        photo_layout = QVBoxLayout(photo_widget)
        self.photo_editor = PhotoEditor()
        self.photo_editor.textRegionSelected.connect(self.on_photo_text_selected)
        photo_layout.addWidget(self.photo_editor)
        self.tabs.addTab(photo_widget, "Photo Editor")
        
        layout.addWidget(self.tabs)
        
        return panel
    
    def open_pdf(self):
        """Open PDF file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open PDF", "", "PDF Files (*.pdf)"
        )
        if file_path:
            self.pdf_path = file_path
            if self.pdf_canvas.load_pdf(file_path):
                self.update_page_label()
                self.statusBar.showMessage(f"Loaded PDF: {os.path.basename(file_path)}")
                self.tabs.setCurrentIndex(0)
    
    def open_photo(self):
        """Open photo file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open Photo", "", 
            "Image Files (*.png *.jpg *.jpeg *.bmp *.tiff *.gif)"
        )
        if file_path:
            self.photo_path = file_path
            if self.photo_editor.load_image(file_path):
                self.statusBar.showMessage(f"Loaded Photo: {os.path.basename(file_path)}")
                self.tabs.setCurrentIndex(1)
                # Auto-extract text with progress
                self.statusBar.showMessage("Extracting text from image...")
                regions = self.photo_editor.extract_text_with_ocr()
                if regions:
                    self.statusBar.showMessage(f"Found {len(regions)} text regions")
                else:
                    self.statusBar.showMessage("No text found or OCR not available")
    
    def toggle_edit_mode(self, checked):
        """Toggle edit mode"""
        self.pdf_canvas.edit_mode = checked
        if checked:
            self.statusBar.showMessage("Edit mode: Click on text to edit")
        else:
            self.statusBar.showMessage("View mode")
    
    def on_text_selected(self, x, y, text, font_info):
        """Handle text selection from PDF"""
        self.text_input.setPlainText(text)
        self.font_combo.setCurrentFont(QFont(font_info.get("font", "Arial")))
        self.font_size.setValue(int(font_info.get("size", 12)))
        
        # Store selection info
        self.selected_position = (x, y)
        self.selected_font_info = font_info
        self.statusBar.showMessage(f"Selected text: {text[:50]}...")
    
    def on_photo_text_selected(self, region):
        """Handle text region selection from photo"""
        self.text_input.setPlainText(region["text"])
        self.font_combo.setCurrentFont(QFont(region["font_info"].get("font", "Arial")))
        self.font_size.setValue(int(region["font_info"].get("size", 12)))
        
        # Store selection info
        self.selected_photo_region = region
        self.statusBar.showMessage(f"Selected text from photo: {region['text'][:50]}...")
    
    def choose_color(self):
        """Open color picker dialog"""
        color = QColorDialog.getColor(self.current_color, self)
        if color.isValid():
            self.current_color = color
            self.color_btn.setStyleSheet(f"background-color: {color.name()}; color: white;")
    
    def apply_text_changes(self):
        """Apply text changes to the document"""
        new_text = self.text_input.toPlainText()
        if not new_text:
            return
        
        # Save state for undo
        self.save_state()
        
        # Get font settings
        font_name = self.font_combo.currentFont().family()
        font_size = self.font_size.value()
        color = self.current_color
        
        # Apply to PDF if in PDF mode
        if self.tabs.currentIndex() == 0 and hasattr(self, 'selected_position'):
            self.apply_pdf_edit(new_text, font_name, font_size, color)
        
        # Apply to photo if in photo mode
        elif self.tabs.currentIndex() == 1:
            if hasattr(self, 'selected_photo_region') and self.selected_photo_region:
                self.apply_photo_edit(new_text, font_name, font_size, color)
            else:
                QMessageBox.information(
                    self, "No Selection", 
                    "Please click on a text region in the photo first, or use OCR to extract text."
                )
                return
        
        self.statusBar.showMessage("Changes applied")
    
    def apply_pdf_edit(self, text, font_name, font_size, color):
        """Apply text edit to PDF with precise positioning"""
        if not self.pdf_canvas.pdf_doc:
            return
        
        page = self.pdf_canvas.pdf_doc[self.pdf_canvas.current_page]
        x, y = self.selected_position
        
        # Get the original text block to get exact dimensions
        original_block = None
        for block in self.pdf_canvas.text_blocks:
            if abs(block["x"] - x) < 5 and abs(block["y"] - y) < 5:
                original_block = block
                break
        
        if original_block:
            # Use original block dimensions for redaction
            rect = fitz.Rect(
                original_block["x"],
                original_block["y"],
                original_block["x"] + original_block["width"],
                original_block["y"] + original_block["height"]
            )
        else:
            # Fallback to approximate size
            rect = fitz.Rect(x, y, x + len(text) * font_size * 0.6, y + font_size * 1.2)
        
        # Delete old text (by redacting)
        page.add_redact_annot(rect)
        page.apply_redactions()
        
        # Insert new text at the same position
        point = fitz.Point(x, y + font_size * 0.8)  # Adjust baseline
        
        # Try to use the font, fallback to helv if not available
        try:
            page.insert_text(
                point,
                text,
                fontsize=font_size,
                fontname=font_name,
                color=(color.red()/255, color.green()/255, color.blue()/255)
            )
        except:
            # Fallback to helv if font not available
            page.insert_text(
                point,
                text,
                fontsize=font_size,
                fontname="helv",
                color=(color.red()/255, color.green()/255, color.blue()/255)
            )
        
        # Refresh display
        self.pdf_canvas.display_page(self.pdf_canvas.current_page)
    
    def apply_photo_edit(self, text, font_name, font_size, color):
        """Apply text edit to photo by overlaying new text"""
        if not self.photo_editor.image:
            return
        
        # Save state for undo
        if not hasattr(self, 'photo_history'):
            self.photo_history = []
        self.photo_history.append(self.photo_editor.image.copy())
        
        # Create a copy to edit
        img = self.photo_editor.image.copy()
        draw = ImageDraw.Draw(img)
        
        # Try to load the font, fallback to default
        try:
            # Try to use system font
            font_path = None
            if sys.platform == 'win32':
                font_path = f"C:/Windows/Fonts/{font_name}.ttf"
            elif sys.platform == 'darwin':
                font_path = f"/Library/Fonts/{font_name}.ttf"
            else:
                font_path = f"/usr/share/fonts/truetype/{font_name}.ttf"
            
            if font_path and os.path.exists(font_path):
                font_obj = ImageFont.truetype(font_path, int(font_size))
            else:
                font_obj = ImageFont.load_default()
        except:
            font_obj = ImageFont.load_default()
        
        # Get selected region if available
        if hasattr(self, 'selected_photo_region') and self.selected_photo_region:
            x, y = self.selected_photo_region["x"], self.selected_photo_region["y"]
        else:
            # Default position (center)
            x, y = img.width // 2, img.height // 2
        
        # Draw text with background for better visibility
        bbox = draw.textbbox((x, y), text, font=font_obj)
        padding = 5
        draw.rectangle(
            [bbox[0] - padding, bbox[1] - padding, bbox[2] + padding, bbox[3] + padding],
            fill=(255, 255, 255, 200)
        )
        
        # Draw the text
        draw.text(
            (x, y),
            text,
            font=font_obj,
            fill=(color.red(), color.green(), color.blue())
        )
        
        # Update image
        self.photo_editor.image = img
        self.photo_editor.display_image()
        self.statusBar.showMessage("Text applied to photo")
    
    def prev_page(self):
        """Go to previous page"""
        if self.pdf_canvas.pdf_doc:
            if self.pdf_canvas.current_page > 0:
                self.pdf_canvas.display_page(self.pdf_canvas.current_page - 1)
                self.update_page_label()
    
    def next_page(self):
        """Go to next page"""
        if self.pdf_canvas.pdf_doc:
            if self.pdf_canvas.current_page < len(self.pdf_canvas.pdf_doc) - 1:
                self.pdf_canvas.display_page(self.pdf_canvas.current_page + 1)
                self.update_page_label()
    
    def update_page_label(self):
        """Update page label"""
        if self.pdf_canvas.pdf_doc:
            total = len(self.pdf_canvas.pdf_doc)
            current = self.pdf_canvas.current_page + 1
            self.page_label.setText(f"Page {current} of {total}")
    
    def on_zoom_changed(self, value):
        """Handle zoom slider change"""
        zoom_percent = value / 100.0
        self.zoom_label.setText(f"{value}%")
        
        # Reset transform and apply new zoom
        self.pdf_canvas.resetTransform()
        self.pdf_canvas.scale(zoom_percent, zoom_percent)
        self.pdf_canvas.zoom_factor = zoom_percent
    
    def save_state(self):
        """Save current state for undo"""
        if self.pdf_canvas.pdf_doc:
            state = {
                "type": "pdf",
                "page": self.pdf_canvas.current_page,
                "doc_data": self.pdf_canvas.pdf_doc.tobytes()
            }
            self.edit_history = self.edit_history[:self.history_index + 1]
            self.edit_history.append(state)
            self.history_index = len(self.edit_history) - 1
    
    def undo(self):
        """Undo last action"""
        if self.history_index > 0:
            self.history_index -= 1
            self.restore_state()
    
    def redo(self):
        """Redo last action"""
        if self.history_index < len(self.edit_history) - 1:
            self.history_index += 1
            self.restore_state()
    
    def restore_state(self):
        """Restore state from history"""
        if 0 <= self.history_index < len(self.edit_history):
            state = self.edit_history[self.history_index]
            if state["type"] == "pdf" and self.pdf_canvas.pdf_doc:
                # Reload PDF from saved state
                doc_data = state["doc_data"]
                self.pdf_canvas.pdf_doc = fitz.open(stream=doc_data, filetype="pdf")
                self.pdf_canvas.display_page(state["page"])
                self.update_page_label()
    
    def export_pdf(self):
        """Export as PDF"""
        if not self.pdf_canvas.pdf_doc:
            QMessageBox.warning(self, "No Document", "Please open a PDF first.")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF", "", "PDF Files (*.pdf)"
        )
        if file_path:
            try:
                self.pdf_canvas.pdf_doc.save(file_path)
                QMessageBox.information(self, "Success", f"PDF saved to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save PDF: {str(e)}")
    
    def export_docx(self):
        """Export as DOCX"""
        if not self.pdf_canvas.pdf_doc:
            QMessageBox.warning(self, "No Document", "Please open a PDF first.")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save DOCX", "", "Word Documents (*.docx)"
        )
        if file_path:
            try:
                doc = Document()
                
                for page_num in range(len(self.pdf_canvas.pdf_doc)):
                    page = self.pdf_canvas.pdf_doc[page_num]
                    text = page.get_text()
                    
                    # Add page heading
                    heading = doc.add_heading(f'Page {page_num + 1}', level=1)
                    
                    # Add text
                    paragraph = doc.add_paragraph(text)
                
                doc.save(file_path)
                QMessageBox.information(self, "Success", f"DOCX saved to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save DOCX: {str(e)}")
    
    def export_png(self):
        """Export current page as PNG"""
        if self.tabs.currentIndex() == 0 and self.pdf_canvas.pdf_doc:
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save PNG", "", "PNG Files (*.png)"
            )
            if file_path:
                try:
                    page = self.pdf_canvas.pdf_doc[self.pdf_canvas.current_page]
                    mat = fitz.Matrix(2, 2)
                    pix = page.get_pixmap(matrix=mat)
                    pix.save(file_path)
                    QMessageBox.information(self, "Success", f"PNG saved to {file_path}")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to save PNG: {str(e)}")
        
        elif self.tabs.currentIndex() == 1 and self.photo_editor.image:
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save PNG", "", "PNG Files (*.png)"
            )
            if file_path:
                try:
                    self.photo_editor.image.save(file_path)
                    QMessageBox.information(self, "Success", f"PNG saved to {file_path}")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to save PNG: {str(e)}")
        else:
            QMessageBox.warning(self, "No Document", "Please open a PDF or photo first.")

def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')  # Modern look
    
    # Apply global stylesheet for better appearance
    app.setStyleSheet("""
        QMainWindow {
            background-color: #f5f5f5;
        }
        QGroupBox {
            font-weight: bold;
            border: 2px solid #cccccc;
            border-radius: 5px;
            margin-top: 10px;
            padding-top: 10px;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 5px;
        }
        QPushButton {
            padding: 6px;
            border-radius: 4px;
            background-color: #e0e0e0;
        }
        QPushButton:hover {
            background-color: #d0d0d0;
        }
        QTextEdit {
            border: 1px solid #cccccc;
            border-radius: 3px;
            padding: 5px;
        }
        QSpinBox, QFontComboBox {
            border: 1px solid #cccccc;
            border-radius: 3px;
            padding: 3px;
        }
    """)
    
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()

